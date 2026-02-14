"""Word 文档解析器"""
from docx import Document as DocxDocument
from typing import Dict, Any
from app.utils.hash import md5_bytes
import zipfile
import xml.etree.ElementTree as ET
import logging
import os

logger = logging.getLogger(__name__)


class DocxParser:
    """DOCX 全量解析引擎"""

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        try:
            doc = DocxDocument(file_path)
            result = {
                "full_text": DocxParser._extract_text(doc),
                "metadata": DocxParser._extract_metadata(file_path),
                "format_info": DocxParser._extract_format(doc),
                "images": DocxParser._extract_images(file_path),
                "page_count": 0,  # docx 无直接页数，可估算
            }
            return result
        except Exception as e:
            logger.error(f"DOCX 解析失败: {file_path}, 错误: {e}")
            return {"full_text": "", "metadata": {}, "format_info": {}, "images": [], "page_count": 0, "error": str(e)}

    @staticmethod
    def _extract_text(doc: DocxDocument) -> str:
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)

    @staticmethod
    def _extract_metadata(file_path: str) -> Dict[str, Any]:
        meta = {}
        try:
            # 从 docProps/core.xml 提取标准元数据
            with zipfile.ZipFile(file_path, 'r') as z:
                if 'docProps/core.xml' in z.namelist():
                    core_xml = z.read('docProps/core.xml')
                    root = ET.fromstring(core_xml)
                    ns = {
                        'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
                        'dc': 'http://purl.org/dc/elements/1.1/',
                        'dcterms': 'http://purl.org/dc/terms/',
                    }
                    for tag, key in [
                        ('dc:creator', 'author'),
                        ('cp:lastModifiedBy', 'last_modified_by'),
                        ('dcterms:created', 'created_date'),
                        ('dcterms:modified', 'modified_date'),
                        ('dc:title', 'title'),
                        ('dc:description', 'description'),
                        ('cp:revision', 'revision'),
                    ]:
                        elem = root.find(tag, ns)
                        meta[key] = elem.text if elem is not None else ""

                # 从 docProps/app.xml 提取应用元数据
                if 'docProps/app.xml' in z.namelist():
                    app_xml = z.read('docProps/app.xml')
                    root = ET.fromstring(app_xml)
                    ns_app = {'ep': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'}
                    for tag, key in [
                        ('ep:Application', 'application'),
                        ('ep:AppVersion', 'app_version'),
                        ('ep:Company', 'company'),
                        ('ep:TotalTime', 'total_editing_time'),
                        ('ep:Template', 'template'),
                    ]:
                        elem = root.find(tag, ns_app)
                        meta[key] = elem.text if elem is not None else ""

        except Exception as e:
            logger.warning(f"元数据提取失败: {e}")

        return meta

    @staticmethod
    def _extract_format(doc: DocxDocument) -> Dict[str, Any]:
        fonts = set()
        font_sizes = set()

        for para in doc.paragraphs[:50]:  # 抽样前50段
            for run in para.runs:
                if run.font.name:
                    fonts.add(run.font.name)
                if run.font.size:
                    font_sizes.add(run.font.size.pt if run.font.size else 0)

        # 页面设置
        section = doc.sections[0] if doc.sections else None
        page_info = {}
        if section:
            page_info = {
                "page_width_mm": round(section.page_width.mm, 1) if section.page_width else 0,
                "page_height_mm": round(section.page_height.mm, 1) if section.page_height else 0,
                "left_margin_mm": round(section.left_margin.mm, 1) if section.left_margin else 0,
                "right_margin_mm": round(section.right_margin.mm, 1) if section.right_margin else 0,
                "top_margin_mm": round(section.top_margin.mm, 1) if section.top_margin else 0,
                "bottom_margin_mm": round(section.bottom_margin.mm, 1) if section.bottom_margin else 0,
            }

        return {
            "fonts": sorted(list(fonts)),
            "font_sizes": sorted(list(font_sizes)),
            **page_info,
        }

    @staticmethod
    def _extract_images(file_path: str) -> list:
        images = []
        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                for name in z.namelist():
                    if name.startswith('word/media/'):
                        data = z.read(name)
                        ext = os.path.splitext(name)[1].lstrip('.')
                        images.append({
                            "name": os.path.basename(name),
                            "size": len(data),
                            "md5": md5_bytes(data),
                            "ext": ext,
                        })
        except Exception as e:
            logger.warning(f"图片提取失败: {e}")
        return images
